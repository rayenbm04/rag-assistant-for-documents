"""
Unit tests for document extraction functions.
Runs without Ollama or ChromaDB (conftest.py stubs them out).

Run with:
    pytest tests/ -v
"""

import os
import sys
import pytest
from unittest.mock import MagicMock

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "venv"))
import main  # noqa: E402


# ──────────────────────────────────────────
# .txt extraction
# ──────────────────────────────────────────

class TestExtractTxt:
    def test_basic_content(self, tmp_path):
        f = tmp_path / "hello.txt"
        f.write_text("Hello, world!\nLine two.", encoding="utf-8")
        result = main.extract_txt_content(str(f), "hello.txt")
        assert "Hello, world!" in result
        assert "Line two." in result

    def test_filename_in_header(self, tmp_path):
        f = tmp_path / "notes.txt"
        f.write_text("Some notes", encoding="utf-8")
        result = main.extract_txt_content(str(f), "notes.txt")
        assert "notes.txt" in result

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        result = main.extract_txt_content(str(f), "empty.txt")
        assert isinstance(result, str)

    def test_utf8_content(self, tmp_path):
        f = tmp_path / "unicode.txt"
        f.write_text("Héllo wörld — こんにちは", encoding="utf-8")
        result = main.extract_txt_content(str(f), "unicode.txt")
        assert "Héllo" in result


# ──────────────────────────────────────────
# .docx extraction
# ──────────────────────────────────────────

class TestExtractDocx:
    def _make_docx(self, tmp_path, paragraphs=None, tables=None):
        from docx import Document
        doc = Document()
        for text in (paragraphs or []):
            doc.add_paragraph(text)
        for rows in (tables or []):
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            for r, row in enumerate(rows):
                for c, cell_text in enumerate(row):
                    table.cell(r, c).text = cell_text
        path = str(tmp_path / "test.docx")
        doc.save(path)
        return path

    def test_paragraphs_extracted(self, tmp_path):
        path = self._make_docx(tmp_path, paragraphs=["First paragraph", "Second paragraph"])
        result = main.extract_docx_content(path, "test.docx")
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_table_extracted(self, tmp_path):
        path = self._make_docx(
            tmp_path,
            tables=[[["Name", "Age"], ["Alice", "30"], ["Bob", "25"]]]
        )
        result = main.extract_docx_content(path, "test.docx")
        assert "Name" in result
        assert "Alice" in result
        assert "Bob" in result

    def test_filename_in_header(self, tmp_path):
        path = self._make_docx(tmp_path, paragraphs=["Content"])
        result = main.extract_docx_content(path, "test.docx")
        assert "test.docx" in result

    def test_empty_document(self, tmp_path):
        path = self._make_docx(tmp_path)
        result = main.extract_docx_content(path, "empty.docx")
        assert isinstance(result, str)


# ──────────────────────────────────────────
# Token tracking
# ──────────────────────────────────────────

class TestRecordTokens:
    def setup_method(self):
        main.token_usage["prompt"] = 0
        main.token_usage["completion"] = 0
        main.token_usage["requests"] = 0

    def test_accumulates_tokens(self):
        mock_result = MagicMock()
        mock_result.raw = {"prompt_eval_count": 100, "eval_count": 50}
        main.record_tokens(mock_result)
        assert main.token_usage["prompt"] == 100
        assert main.token_usage["completion"] == 50

    def test_accumulates_across_calls(self):
        mock = MagicMock()
        mock.raw = {"prompt_eval_count": 10, "eval_count": 5}
        main.record_tokens(mock)
        main.record_tokens(mock)
        assert main.token_usage["prompt"] == 20
        assert main.token_usage["completion"] == 10

    def test_missing_raw_does_not_crash(self):
        mock = MagicMock()
        mock.raw = None
        main.record_tokens(mock)
        assert main.token_usage["prompt"] == 0

    def test_returns_token_counts(self):
        mock = MagicMock()
        mock.raw = {"prompt_eval_count": 42, "eval_count": 17}
        p, c = main.record_tokens(mock)
        assert p == 42
        assert c == 17
